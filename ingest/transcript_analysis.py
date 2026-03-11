import os
import re
import json
import sqlite3
import argparse
from typing import Dict, List, Any
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

# Note: In a real environment, we would use an LLM client here.
# For this script, we will simulate the extraction for the demo or structure it for future use.

class TranscriptAnalyzer:
    def __init__(self, prompt_template_path: str):
        with open(prompt_template_path, 'r') as f:
            self.template = f.read()

    def load_transcript(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content

    def extract_metadata(self, file_path: str) -> Dict[str, str]:
        filename = os.path.basename(file_path)
        # Try to find date YYYY-MM-DD
        date_match = re.search(r'(\d{4}-\d{2}-\d{2})', filename)
        date = date_match.group(1) if date_match else "Unknown"
        
        return {
            "memo_id": f"MEMO-{date}-01",
            "date_of_meeting": date,
            "subject": filename.replace('.txt', '').replace('_', ' ')
        }

    def run_analysis(self, transcript_content: str, metadata: Dict[str, str]) -> Dict[str, Any]:
        """
        Placeholder for LLM analysis. In production, this would call Ollama.
        """
        memo = {
            "memo_id": metadata["memo_id"],
            "subject": metadata["subject"],
            "date_of_meeting": metadata["date_of_meeting"],
            "speakers": ["Unknown"],
            "summary": "Meeting transcript analysis placeholder.",
            "key_findings": [
                {
                    "category": "Admission",
                    "description": "Discussed favor flex move payments.",
                    "timestamp": "[00:05:30.00]",
                    "quote": "Can we pay him for a favor flex move?",
                    "legal_significance": "Admits to potential off-books payment structures.",
                    "suggested_action": "Cross-reference with bank statements."
                }
            ],
            "resolved_discrepancies": [],
            "unresolved_questions": ["What is a 'favor flex move'?"],
            "exhibits_referenced": []
        }
        return memo

    def find_related_evidence(self, db_path: Path, meeting_date: str, window_days: int = 7) -> List[Dict[str, Any]]:
        """
        Queries evidence_hub.db for records within a window of the meeting date.
        """
        if meeting_date == "Unknown":
            return []
            
        try:
            meeting_dt = datetime.strptime(meeting_date, "%Y-%m-%d")
        except ValueError:
            return []

        start_dt = (meeting_dt - timedelta(days=window_days)).isoformat()
        end_dt = (meeting_dt + timedelta(days=window_days)).isoformat()

        conn = sqlite3.connect(str(db_path))
        conn.row_factory = sqlite3.Row
        
        query = """
            SELECT source_type, title, summary, start_timestamp, body_snippet, tags
            FROM evidence
            WHERE start_timestamp BETWEEN ? AND ?
            ORDER BY start_timestamp ASC
            LIMIT 50
        """
        
        related = []
        try:
            cur = conn.cursor()
            rows = cur.execute(query, (start_dt, end_dt)).fetchall()
            for row in rows:
                related.append(dict(row))
        finally:
            conn.close()
            
        return related

    def generate_docx_memo(self, memo: Dict[str, Any], related_evidence: List[Dict[str, Any]], output_path: str):
        doc = Document()
        
        # Header
        doc.add_heading('PARALEGAL TRANSCRIPT MEMO', 0)
        
        # Metadata Table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        def add_row(t, k, v):
            cells = t.rows[len(t.rows)-1].cells
            if len(t.rows) > 1 and not cells[0].text:
                row = cells
            else:
                row = t.add_row().cells
            row[0].text = k
            row[1].text = str(v)

        # Pre-populate first 4 rows
        table.cell(0, 0).text = "MEMO ID"
        table.cell(0, 1).text = memo.get("memo_id", "")
        table.cell(1, 0).text = "SUBJECT"
        table.cell(1, 1).text = memo.get("subject", "")
        table.cell(2, 0).text = "DATE"
        table.cell(2, 1).text = memo.get("date_of_meeting", "")
        table.cell(3, 0).text = "SPEAKERS"
        table.cell(3, 1).text = ", ".join(memo.get("speakers", []))

        # Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(memo.get("summary", ""))

        # Key Findings
        doc.add_heading('Key Findings & Legal Significance', level=1)
        for finding in memo.get("key_findings", []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{finding.get('category')}] ").bold = True
            p.add_run(f"Timestamp: {finding.get('timestamp')}\n")
            p.add_run(f"Quote: \"{finding.get('quote')}\"\n")
            p.add_run(f"Significance: {finding.get('legal_significance')}")

        # Related Evidence
        if related_evidence:
            doc.add_heading('Cross-Referenced Evidence (Evidence Hub)', level=1)
            for ev in related_evidence:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"[{ev.get('source_type').upper()}] ").bold = True
                p.add_run(f"{ev.get('title')}\n")
                p.add_run(f"Date: {ev.get('start_timestamp')}\n")
                p.add_run(f"Summary: {ev.get('summary')}")

        doc.save(output_path)
        print(f"✅ Memo saved to {output_path}")

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    default_template = os.path.join(project_root, 'prompts', 'forensic_analysis.txt')

    parser = argparse.ArgumentParser(description='Analyze legal transcripts.')
    parser.add_argument('file', help='Path to the transcript file')
    parser.add_argument('--template', default=default_template, help='Path to the prompt template')
    parser.add_argument('--output', help='Path to save the JSON memo')
    parser.add_argument('--docx', help='Path to save the DOCX memo')
    parser.add_argument('--cross-ref', action='store_true', help='Cross-reference with Evidence Hub')
    parser.add_argument('--window', type=int, default=7, help='Window in days for cross-referencing')
    parser.add_argument('--db', default=os.path.join(project_root, 'data', 'evidence_hub.db'), help='Path to evidence_hub.db')

    args = parser.parse_args()
    
    analyzer = TranscriptAnalyzer(args.template)
    content = analyzer.load_transcript(args.file)
    metadata = analyzer.extract_metadata(args.file)
    
    print(f"🚀 Analyzing {args.file}...")
    memo = analyzer.run_analysis(content, metadata)
    
    related = []
    if args.cross_ref:
        print(f"🔍 Cross-referencing with Evidence Hub (Window: {args.window} days)...")
        related = analyzer.find_related_evidence(Path(args.db), metadata["date_of_meeting"], args.window)
        print(f"✅ Found {len(related)} related records.")

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(memo, f, indent=2)
        print(f"✅ JSON saved to {args.output}")

    if args.docx:
        output_path = args.docx
        analyzer.generate_docx_memo(memo, related, output_path)

if __name__ == "__main__":
    main()
